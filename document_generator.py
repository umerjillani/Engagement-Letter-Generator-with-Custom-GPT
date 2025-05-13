import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING,WD_TAB_ALIGNMENT
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, Cm

import json
import os
import re
from datetime import datetime

def create_engagement_letter(template_json_path, user_data_json_path, output_path):
    """
    Creates a formatted DOCX engagement letter based on template JSON and user data.
    
    Args:
        template_json_path (str): Path to the structured template JSON file
        user_data_json_path (str): Path to the user data JSON file
        output_path (str): Path where the output DOCX file will be saved
    """
    # Load JSON data
    with open(template_json_path, 'r') as template_file:
        template_data = json.load(template_file)
    
    with open(user_data_json_path, 'r') as user_file:
        user_data = json.load(user_file)
    
    # Create a new document
    doc = docx.Document()
    
    # Set default font to Calibri for the whole document
    set_default_font(doc, 'Calibri')
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add header with logo (placeholder for actual logo)
    first_section = doc.sections[0]
    first_section.different_first_page_header_footer = True

    # --- FIRST PAGE HEADER ---
    first_page_header = first_section.first_page_header

    # Create a table for the header to align logo and company details side by side
    # Create a table with proper width parameters
    header_table = first_page_header.add_table(rows=1, cols=2, width=Inches(6.0))
    header_table.autofit = True
    
    # Left cell for the logo
    logo_cell = header_table.cell(0, 0)
    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_run = logo_para.add_run()
    logo_run.add_picture(
        template_data['header']['logo_ref'], 
        width=Inches(2.08)  # Adjust width as needed
    )

    # Right cell for company details
    details_cell = header_table.cell(0, 1)
    details_para = details_cell.paragraphs[0]
    details_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add the details
    details_para.add_run(template_data['header']['company_details']['phone']).bold = True
    details_para.add_run('\n' + template_data['header']['company_details']['address'])
    details_para.add_run('\n' + template_data['header']['company_details']['city_state_zip'])

    # --- GENERAL FOOTER (for non-first pages) ---
    general_footer = first_section.footer
    footer_para = general_footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add footer image
    footer_run = footer_para.add_run()
    footer_run.add_picture(
        template_data['footer']['image_ref'],
        width=Inches(2.89)  # Adjust width as needed
    )

    # Position footer with appropriate spacing (must be positive)
    footer_para.paragraph_format.space_before = Pt(12)    
    # Add date
    date_para = doc.add_paragraph()
    date_para.add_run(user_data['header']['date'])
    date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add client info
    client_para = doc.add_paragraph()
    client_para.add_run(user_data['header']['client_name']).bold = True
    
    # TITLE (Optional)
    if user_data['header'].get('title'):
        client_para.add_run('\n' + user_data['header']['title'])
    client_para.add_run('\n' + user_data['header']['company_name'])
    
    # Add address lines
    for address_line in user_data['header']['address_lines']:
        client_para.add_run('\n' + address_line)
    
    # 5. CONTACT NUMBER (Optional)
    if user_data['header'].get('contact_number'):
        client_para.add_run('\n' + user_data['header']['contact_number'])
    client_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add subject line
    subject_para = doc.add_paragraph()
    subject_para.add_run(user_data['subject_line']).bold = True
    subject_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add salutation
    salutation_para = doc.add_paragraph()
    salutation_para.add_run(user_data['salutation'])
    salutation_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add introduction section
    intro_para = doc.add_paragraph()
    intro_text= template_data['sections']['introduction']['description']
    
    intro_text_replaced = intro_text.replace("{header.client_name}", user_data['header']['client_name'])
    intro_text_replaced = intro_text_replaced.replace("{header.company_name}", user_data['header']['company_name'])

    intro_para.add_run(intro_text_replaced)

    set_justified_paragraph(intro_para)
    
    # Add Description of Services section with main heading
    add_main_heading(doc, "I. Description of Services")
    
    # Add Service Description subsections
    add_service_description(doc, user_data['description_of_services'])
    
    # Add Services Disclaimer section with main heading
    add_main_heading(doc, "III. Services Disclaimer")
    
    for content_item in template_data['sections']['services_disclaimer']['content']:
        disclaimer_para = doc.add_paragraph()
        disclaimer_para.add_run(content_item['content'])
        set_justified_paragraph(disclaimer_para)
    
    # Add Client Responsibilities section with main heading
    add_main_heading(doc, "IV. CLIENT Responsibilities and Information")
    
    for content_item in template_data['client_responsibilities_and_information']['content']:
        resp_para = doc.add_paragraph()
        for inner_content in content_item['content']:
            text = inner_content['content']
            if is_bulleted_text(text):
                add_bulleted_list(doc, text)
            else:
                resp_para.add_run(text)
                set_justified_paragraph(resp_para)
    
    # Add Fees and Expenses section with main heading
    add_main_heading(doc, "II. Fees And Expenses")
    
    fees_intro_para = doc.add_paragraph()
    fees_intro_para.add_run(user_data['fees_and_expenses']['start_description'])
    set_justified_paragraph(fees_intro_para)
    
    # Add Fee subsections
    add_fees_sections(doc, user_data['fees_and_expenses'])
    
    # Add remaining sections - pass user_data to the function
    add_remaining_sections(doc, template_data, user_data)
    
    # Add signature block
    doc.add_paragraph()
    
    # Add signature content with right alignment - no table
    for content_item in template_data['signature']['content']:
        sig_para = doc.add_paragraph()
        sig_para.add_run(content_item['content'])
        sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
    # Replace client name in signature block with correct name
    replace_text_in_document(doc, "{header.client_name}", user_data['header']['client_name'])
    replace_text_in_document(doc, "{header.title}", user_data['header']['title'])
    
    # Replace other dynamic fields
    replace_text_in_document(doc, "{header.company_name}", user_data['header']['company_name'])
    replace_text_in_document(doc, "{header.address_lines}", "\n".join(user_data['header']['address_lines']))
    replace_text_in_document(doc, "{header.client_email}", user_data['header']['client_email'])
    
    # Save the document
    doc.save(output_path)
    print(f"Engagement letter successfully created at {output_path}")

def set_default_font(doc, font_name):
    """Set default font for the entire document"""
    # Access the styles
    styles = doc.styles
    
    # Modify the Normal style to use Calibri
    normal_style = styles['Normal']
    font = normal_style.font
    font.name = font_name
    font.size = Pt(11)
    
    # Also update other commonly used styles
    for style_name in ['Heading 1', 'Heading 2', 'Heading 3', 'List Bullet', 'List Number']:
        if style_name in styles:
            style = styles[style_name]
            style.font.name = font_name

def add_main_heading(doc, text):
    """Add a main heading with Roman numeral style"""
    heading_para = doc.add_paragraph()
    run = heading_para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    heading_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_subheading(doc, text, letter=None):
    """Add a subheading with appropriate formatting, with content on same line"""
    heading_para = doc.add_paragraph()
    
    if letter:
        heading_text = f"{letter}. {text}"
        run = heading_para.add_run(heading_text)
        run.underline = True
        run.font.size = Pt(11.5)
    else:
        run = heading_para.add_run(text)
        run.underline = True
        run.font.size = Pt(11.5)
        
    heading_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Return the paragraph so content can be added to the same line
    return heading_para

def set_justified_paragraph(paragraph):
    """Set paragraph to justified alignment with appropriate spacing"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    # Set first line indent to help with justified text appearance
    paragraph.paragraph_format.first_line_indent = Pt(12)

def is_bulleted_text(text):
    """
    Detect if text contains bullet points or line items.
    This function checks for common bullet point patterns.
    """
    # Check for patterns that indicate this might be a bulleted list:
    bullet_indicators = [
        # Text has lines starting with bullet-like characters
        bool(re.search(r'(?m)^[\s]*[•\-\*\+\u2022\u2023\u25E6\u2043\u2219][\s]+', text)),
        
        # Text has lines starting with numbers or letters followed by periods or parentheses
        bool(re.search(r'(?m)^[\s]*[a-zA-Z0-9]+[\.\)][\s]+', text)),
        
        # Text contains semicolons at end of lines and starts with capital letters
        bool(re.search(r';\s*\n[A-Z]', text)),
        
        # Check for lines starting with action words with similar patterns
        bool(re.search(r'(?m)^[\s]*(Review|Analyze|Prepare|Assist|Research|Draft|When|Provide|Create|Develop)', text) and 
             (text.count('\n') > 2 or text.count(';') > 2)),
        
        # If many lines start with a similar pattern (typical for lists)
        len(set(re.findall(r'(?m)^[\s]*(\w+)', text))) < text.count('\n') / 2 and text.count('\n') > 3,
        
        # Check if text has multiple lines with capital starts and all ending with either ; or .
        text.count('\n') > 2 and all(line.strip().endswith(('.', ';')) for line in text.split('\n') if line.strip())
    ]
    
    # If text mentions specific financial roles or hourly rates - likely a list
    financial_list = (
        ("Fiduciary" in text and "Director" in text and "$" in text) or
        ("hourly" in text.lower() and "rate" in text.lower() and "$" in text)
    )
    
    # Return True if any bullet indicator is True or if it's a financial list
    return any(bullet_indicators) or financial_list

def add_bulleted_list(doc, text):
    """
    Convert text to a properly formatted bulleted list.
    """
    # Split text into lines
    lines = text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Remove common bullet prefixes
        cleaned_line = re.sub(r'^[\s]*[•\-\*\+\u2022\u2023\u25E6\u2043\u2219][\s]+', '', line)
        
        # Remove alphabetic/numeric bullets (like "a." or "1.")
        cleaned_line = re.sub(r'^[\s]*[a-zA-Z0-9]+[\.\)][\s]+', '', cleaned_line)
        
        # Apply standard bullet style
        bullet_para = doc.add_paragraph(style='List Bullet')
        bullet_para.add_run(cleaned_line)
        bullet_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        bullet_para.paragraph_format.space_after = Pt(6)
        bullet_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

def add_service_description(doc, description_data):
    """Add service description sections to the document"""
    
    # Start description
    if "start_description" in description_data:
        start_para = doc.add_paragraph()
        start_para.add_run(description_data['start_description'])
        set_justified_paragraph(start_para)
    
    # Add subsections
    for i, subsection in enumerate(description_data['subsections']):
        if subsection['title']:
            # Create paragraph for subsection heading
            subsection_para = add_subheading(doc, subsection['title'], letter=chr(ord('A')+i))
            
            # Add first paragraph of content in the same line if available
            if 'content' in subsection and subsection['content']:
                first_content = subsection['content'][0]
                if first_content['type'] == 'paragraph':
                    # Add a space after the heading and add the content
                    subsection_para.add_run(" ").underline = False
                    subsection_para.add_run(first_content['text'])
                    set_justified_paragraph(subsection_para)
                    # Remove this item as it's already added
                    subsection['content'] = subsection['content'][1:]
        
        # Handle nested subsections
        if 'subsections' in subsection:
            for j, nested_subsection in enumerate(subsection['subsections'], 1):
                nested_title_para = doc.add_paragraph()
                nested_title_para.add_run(f"{j}. {nested_subsection['title']}")
                nested_title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Handle content in nested subsections
                if 'content' in nested_subsection:
                    for content_item in nested_subsection['content']:
                        if content_item['type'] == 'bullet':
                            # This is explicitly marked as a bullet list
                            for item in content_item['items']:
                                bullet_para = doc.add_paragraph(style='List Bullet')
                                bullet_para.add_run(item.strip())
                                bullet_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                bullet_para.paragraph_format.space_after = Pt(6)
                        else:
                            # Check if this might contain bullets
                            text = content_item.get('text', '')
                            if is_bulleted_text(text):
                                add_bulleted_list(doc, text)
                            else:
                                para = doc.add_paragraph()
                                para.add_run(text)
                                set_justified_paragraph(para)
        
        # Handle remaining direct content
        if 'content' in subsection:
            for content_item in subsection['content']:
                if content_item['type'] == 'paragraph':
                    text = content_item['text']
                    if is_bulleted_text(text):
                        add_bulleted_list(doc, text)
                    else:
                        para = doc.add_paragraph()
                        para.add_run(text)
                        set_justified_paragraph(para)

def add_fees_sections(doc, fees_data):
    """Add fees and expenses sections to the document"""
    
    # Add subsections
    for i, subsection in enumerate(fees_data['subsections']):
        letter = chr(ord('A')+i)
        
        # Create paragraph for subsection heading
        subsection_para = add_subheading(doc, subsection['title'], letter=letter)
        
        # Add first content item on the same line if available
        if 'content' in subsection and subsection['content']:
            first_content = subsection['content'][0]
            # Add a space after the heading and add the content
            subsection_para.add_run(" ").underline = False
            subsection_para.add_run(first_content['text'])
            set_justified_paragraph(subsection_para)
            # Skip this item in the later loop
            remaining_content = subsection['content'][1:]
        else:
            remaining_content = subsection['content']
        
        # Add remaining content
        for content_item in remaining_content:
            text = content_item['text']
            
            # Check if this is a rate/fee listing or other bulleted content
            if is_bulleted_text(text):
                add_bulleted_list(doc, text)
            else:
                para = doc.add_paragraph()
                para.add_run(text)
                set_justified_paragraph(para)
        
        # Add end description if present
        if 'end_description' in subsection:
            end_para = doc.add_paragraph()
            end_para.add_run(subsection['end_description'])
            set_justified_paragraph(end_para)

def add_remaining_sections(doc, template_data, user_data):
    """Add all remaining sections from the template"""
    
    section_order = [
        ("V. Non-solicitation", "non_solicitation"),
        ("VI. Term and Termination", "term_and_termination"),
        ("VII. No Third-Party Beneficiary", "no_third_party_beneficiary"),
        ("VIII. Conflicts", "conflicts"),
        ("IX. Confidentiality / Non-Solicitation.", "confidentiality_non_solicitation"),
        ("X. Indemnification.", "Indemnification"),
        ("XI. Notices.", "notices"),
        ("XII. Limitation of Liability.", "limitation_of_liability"),
        ("XIII. Attorneys' Fees.", "attorneys_fees"),
        ("XIV. Miscellaneous.", "miscellaneous")
    ]
    
    for title, key in section_order:
        add_main_heading(doc, title)
        
        # Special handling for notices section
        if key == "notices":
            # Add the introductory paragraph
            intro_para = doc.add_paragraph()
            intro_para.add_run(template_data[key]['content'][0]['content'])
            set_justified_paragraph(intro_para)
            
            # Handle RESOLUTE address - left aligned
            resolute_address = """If to RESOLUTE:
Resolute Commercial Services, LLC
6750 East Camelback Road, Suite 103
Scottsdale, AZ 85251
Attn: Eric Anderes
Or
Eanderes@resolutecommercial.com"""
            
            lines = resolute_address.split('\n')
            for line in lines:
                if line.strip():
                    line_para = doc.add_paragraph()
                    line_para.add_run(line.strip())
                    line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    line_para.paragraph_format.space_after = Pt(0)
            
            # Add a little space between addresses
            doc.add_paragraph()
            
            # Handle CLIENT address - left aligned
            client_address = f"""if to CLIENT:
{user_data['header']['company_name']}
{' '.join(user_data['header']['address_lines'])}
Attn: {user_data['header']['client_name']}
Or
{user_data['header']['client_email']}"""
            
            lines = client_address.split('\n')
            for line in lines:
                if line.strip():
                    line_para = doc.add_paragraph()
                    line_para.add_run(line.strip())
                    line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    line_para.paragraph_format.space_after = Pt(0)
            
            # Add the final paragraph
            doc.add_paragraph()
            final_para = doc.add_paragraph()
            final_para.add_run(template_data[key]['content'][-1]['content'])
            set_justified_paragraph(final_para)
        elif key == "Indemnification":
            # Special handling for indemnification with subsections
            for i, subsection in enumerate(template_data[key]['content']):
                letter = chr(ord('A')+i)
                
                # Create paragraph for subsection heading
                subsection_para = add_subheading(doc, subsection['title'], letter=letter)
                
                # Add first content item on the same line if available
                if subsection['content']:
                    # Add a space after the heading and add the content
                    subsection_para.add_run(" ").underline = False
                    subsection_para.add_run(subsection['content'][0]['content'])
                    set_justified_paragraph(subsection_para)
                    # Skip this item in the later loop
                    remaining_content = subsection['content'][1:]
                else:
                    remaining_content = subsection['content']
                
                # Add remaining content
                for content_item in remaining_content:
                    text = content_item['content']
                    if is_bulleted_text(text):
                        add_bulleted_list(doc, text)
                    else:
                        para = doc.add_paragraph()
                        para.add_run(text)
                        set_justified_paragraph(para)
        elif key == "miscellaneous":
            # Special handling for miscellaneous with subsections
            for i, subsection in enumerate(template_data[key]['content']):
                letter = chr(ord('A') + i)
                
                # Create paragraph for subsection heading
                subsection_para = add_subheading(doc, subsection['title'], letter=letter)
                
                # Add first content item on the same line if available
                if subsection['content']:
                    # Add a space after the heading and add the content
                    subsection_para.add_run(" ").underline = False
                    subsection_para.add_run(subsection['content'][0]['content'])
                    set_justified_paragraph(subsection_para)
                    # Skip this item in the later loop
                    remaining_content = subsection['content'][1:]
                else:
                    remaining_content = subsection['content']
                
                # Add remaining content
                for content_item in remaining_content:
                    text = content_item['content']
                    if is_bulleted_text(text):
                        add_bulleted_list(doc, text)
                    else:
                        para = doc.add_paragraph()
                        para.add_run(text)
                        set_justified_paragraph(para)
        else:
            # Standard section handling
            content_items = template_data[key]['content']
            first_item = True
            
            for item in content_items:
                if item['type'] == 'paragraph':
                    text = item['content']
                    if is_bulleted_text(text):
                        add_bulleted_list(doc, text)
                    else:
                        para = doc.add_paragraph()
                        para.add_run(text)
                        set_justified_paragraph(para)
                elif item['type'] == 'subsection':
                    if item['title']:
                        # Create paragraph for subsection heading
                        subsection_para = add_subheading(doc, item['title'])
                        
                        # Add first content item on the same line if available
                        if item['content']:
                            # Add a space after the heading and add the content
                            subsection_para.add_run(" ").underline = False
                            subsection_para.add_run(item['content'][0]['content'])
                            set_justified_paragraph(subsection_para)
                            # Skip this item in the later loop
                            remaining_content = item['content'][1:]
                        else:
                            remaining_content = item['content']
                        
                        # Add remaining content
                        for subcontent in remaining_content:
                            text = subcontent['content']
                            if is_bulleted_text(text):
                                add_bulleted_list(doc, text)
                            else:
                                para = doc.add_paragraph()
                                para.add_run(text)
                                set_justified_paragraph(para)
    
    # Add the end description from miscellaneous if present
    if 'end_description' in template_data['miscellaneous']:
        end_para = doc.add_paragraph()
        end_para.add_run(template_data['miscellaneous']['end_description'])
        set_justified_paragraph(end_para)

def replace_text_in_document(doc, placeholder, replacement):
    """Replace all instances of placeholder text in the document"""
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace(placeholder, replacement)
    
    # Also check tables if any
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        for run in paragraph.runs:
                            run.text = run.text.replace(placeholder, replacement)

# Example usage
if __name__ == "__main__":
    template_path = "structured_document.json"
    user_data_path = "userdata.json"
    output_path = "engagement_letter.docx"
    
    create_engagement_letter(template_path, user_data_path, output_path)